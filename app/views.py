# Обработчик ссылок
from calendar import month
from email import message
from operator import truth
from unicodedata import name
from flask import  render_template, flash, request, redirect, url_for,jsonify
from requests import patch
from sqlalchemy import values  # для работы с интернетом
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import login_required, login_user, current_user, logout_user
from app import application, models, db, forms
from app.email import send_password_reset_email
import yadisk
from docx import Document # вставка таблицы в word 
import os
from datetime import datetime, timedelta,date
from app import yToken
from config import Config
from app import functions
from werkzeug.utils import secure_filename
import random
import json
from flask_json import FlaskJSON, JsonError, json_response, as_json



now = datetime.now()# заппуск считывания времени
       
@application.route('/createDb')  # вход
@login_required  # только зарегистрированный человек сможет зайти
def createDb():
    db.create_all()  # создаем БД
    user = models.User(email="root", who=3)
    user.set_password("root")
    db.session.add(user)
    db.session.commit()
        

@application.route('/', methods=['POST', 'GET'])  # вход
def index():
    if current_user.is_authenticated:
        return redirect(url_for('cabinet'))
    form = forms.LoginForm()
    if form.validate_on_submit():
        user = db.session.query(models.User).filter(models.User.email == form.email.data).first()
        print(user)
        if user and user.check_password(form.password.data):
            login_user(user, remember=form.remember.data) # запуск сессии пользователя
            return redirect(url_for('cabinet'))
        return redirect(url_for('index'))
    return render_template('index.html', form=form)


@application.route('/register', methods=['POST', 'GET'])# регистрация
def register():
    if current_user.is_authenticated:
        return redirect(url_for('cabinet'))
    form = forms.RegistrationForm()
    if form.validate_on_submit():
        user = models.User(email=form.email.data, who=3)# блокируем профиль
        user.set_password(form.password.data)
        db.session.add(user)
        userId = db.session.query(models.User).filter(models.User.email == form.email.data).first()
        userData = models.Users_Data(idUser=userId.id, name="none",
                                     birthDAy="none", passport="none", passportData="none",
                                     passportBy="none", passportCod="none", nickname="none",
                                     link_vk="none", inn="none", bank_details="none", bankName="none",
                                     address = "none", bankAccount = "none", tags = "none", ogrn = "none",
                                     up_ooo = "none",phone_number="none",chekManager = "none",avatar ="static/x_06eb1977.jpg")
        db.session.add(userData)
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('register.html', form=form)


@application.route('/logout')  # выход из аккаунта
def logout():
    logout_user()
    return redirect(url_for('index'))


@application.route('/reset_password_request', methods=['GET', 'POST'])  # Страница просьбы о смене пароля
def reset_password_request():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = forms.ResetPasswordRequestForm()
    if form.validate_on_submit():
        user = models.User.query.filter_by(email=form.email.data).first()
        if user:
            send_password_reset_email(user)
        return redirect(url_for('index'))
    return render_template('reset_password_request.html', form=form)


@application.route('/reset_password/<token>', methods=['GET', 'POST'])  # Страница сброса пароля и обработка его смены
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    user = models.User.verify_reset_password_token(token)
    if not user:
        return redirect(url_for('index'))
    form = forms.ResetPasswordForm()
    if form.validate_on_submit():
        user.set_password(form.password.data)
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('reset_password.html', form=form)


@application.route('/cabinet', methods=['POST', 'GET'])  # Личный кабинет
@login_required  # только зарегистрированный человек сможет зайти
def cabinet():
    #print(current_user)# просмотрт под каким пользователем вошел человек
    info = functions.chekCabinet(current_user)# проверка заполненного профиля
    return render_template('cabinet.html', current_user=current_user,info=info)
    


@application.route('/cabinet_changer',methods=['POST', 'GET'])  # Страница для изменения данных в личном кабинете и вноса изменений в базу данных
@login_required  # только зарегистрированный человек сможет зайти
def cabinet_changer():
    form = forms.PersonalForm()
    form.nickname.default = "Саня"
    error = ""
    user_data=db.session.query(models.Users_Data).filter_by(idUser=current_user.id).one()
    if request.method == 'POST':
        functions.info_changer(user_data)
        return redirect(url_for('cabinet'))
    return render_template('cabinet_changer.html', form=form, error = error)


@application.route('/my_tasks', methods=['GET', 'POST'])  # Страница с задачами пользователя
@login_required  # только зарегистрированный человек сможет зайти
def my_tasks(): 
    tasks = models.Tasks.query.filter_by(idUser = current_user.id).all()# таблица задачь
    massege = models.Massege.query.filter_by().all()# таблица сообщений


    if request.method == "POST":
        listForm = request.form.to_dict()# получение информации из формы
        if(listForm.get('tascChek')):
            print(request.files.getlist("file"))
            if (request.files.getlist("file")):
                qeru = models.Tasks.query.filter_by(id = listForm['tascChek']).first()#поиск задачи
                files = request.files.getlist("file")

                for file in files:
                    if file.filename == '':
                        return redirect(request.url)
                    print("я работаю")
                    file.save(os.path.join(Config.way_task, file.filename))
                    #os.remove(way + file.filename)
                functions.recursive_upload(yToken, Config.way_task, qeru.linkDisk +"/")# яндексдиск, откуда, куда
                for file in files:
                    if file.filename == '':
                        return redirect(request.url)
                    print("я удаляю")
                    #file.save(os.path.join(Config.way_task, file.filename))
                    os.remove(Config.way_task + file.filename)

                task = models.Tasks.query.filter_by(id = listForm['tascChek']).one()
                task.statusСompleted = "Сheck"
                db.session.commit()# создание соеденения
                db.session.add(task)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд

        # чат
        if (listForm.get('masseg')):
            print("записываю сообщение")
            user = models.Users_Data.query.filter_by(idUser = current_user.id).one() #User
            masg = models.Massege(id_progect = listForm['id_proj'],id_folder =listForm['id_folder'], 
            id_task = listForm['id_task'],id_user = listForm['id_user'], text = listForm['text'],img = user.avatar)
            db.session.commit()# создание соеденения
            db.session.add(masg)# запись в базу данных users_projekt
            db.session.commit()# закрытие соеденения с бд
            massege = models.Massege.query.filter_by(id_progect = listForm['id_proj'],id_folder = listForm['id_folder']).all()# таблица сообщений
            get = {'success': 'true'}
            dic1 = {}
            for u in massege:
                dic = u.__dict__
                if '_sa_instance_state' in dic.keys():
                    del dic['_sa_instance_state']
                name = 'mas'+ str(dic["id"])
                dic1[name] = dic
                get.update(dic1)
            #print(get)
            return json.dumps(get)
    for u in tasks:
        yToken.publish(u.linkDisk)
        u.linkDisk = yToken.get_meta(u.linkDisk).public_url
    return render_template("my_tasks.html", tasks=tasks, now=now, current_user = current_user,massege=massege)


@application.route('/my_documents', methods=['GET', 'POST'])  # Страница с документами пользователя
@login_required  # только зарегистрированный человек сможет зайти
def my_documents():
    user = models.Users_Data.query.filter_by(idUser = current_user.id).one() #User
    #chek = models.Check.query.filter_by(id = current_user.id ).all()
    List = {}
    if request.method == 'POST':
        listForm = request.form.to_dict()
        listFormKeys = listForm.keys()# выводим все ключи выбранных блоках
        # вставка данных
        for id in listFormKeys:
            user = models.Users_Data.query.filter_by(id = id).one()#поиск юзера
            #print("привет я начал работать с яндекс диском")
            listForm = request.form.to_dict()# получение информации из формы
            #print(user.name)
        files = request.files.getlist("file")
        for file in files:
            print("я работаю")
            file.save(os.path.join(Config.way_check, file.filename))
            #chekAdd = models.Chek(idUser = current_user.id,name = file.filename,idDoc = )
            #db.session.commit()# создание соеденения
            #db.session.add(chekAdd)# запись в базу данных users_projekt
            #db.session.commit()# закрытие соеденения с бд 
            #os.remove(Config.way_check + file.filename)
        functions.recursive_upload(yToken, "app/" + Config.way_check, "/чеки/" + user.name)
        for file in files:
            print("я работаю")
            #file.save(os.path.join(Config.way_check, file.filename))
            os.remove(Config.way_check + file.filename)
            
        return render_template('my_documents.html',list=List, error = "Чек прикреплён и внесен на яндекс диск")

    try:
        docList = list(yToken.listdir('/договора/' + user.name))
        a = 0
        for i in docList:
            d = {docList[a].name : docList[a].file}
            List.update(d)
            a = a + 1  
    except yadisk.exceptions.PathNotFoundError:
        return render_template("my_documents.html",list = List,error = "Договор с вами еще не составлен")
    print(docList)
    return render_template("my_documents.html", list = List,user=user)


base = {'idProject':"", 'id_folder': ""}
@application.route('/my_projects', methods=['GET', 'POST'])  # Страница с проектами шоураннера
@login_required  # только зарегистрированный человек сможет зайти
def my_projects():
    if current_user.who==1 or current_user.who==32:
        error = ""
        projects = models.Project.query.all() #запрос в базу данны для вывода проектов
        user_project = models.Users_Projects.query.all()#запрос в базу данны для вывода какой человек к каому проекту
        for i in user_project:
            print(i.up)
        info = models.User.query.all() #запрос в базу данны для вывода людей
        tasks = models.Tasks.query.all()# таблица задачь
        folders = models.Folder.query.all()# таблица папок
        
        if request.method == 'POST': # обработка post
            listForm = request.form.to_dict()# получение информации из формы
            # добавить людей
            if (listForm.get('thePiple')):
                print("добавить человека------------------------")
                for k, v in listForm.items():
                    if v == 'id_sel':
                        print(k)
                        if (models.Users_Projects.query.filter_by(User_id = k,Project_id = listForm.get('thePiple')).first()):
                            continue
                        else:
                            print("я добовляю:" + k)
                            update_user_projects = models.Users_Projects(User_id = k, Project_id = listForm.get('thePiple'))# обновления таблицы Users_project
                            db.session.add(update_user_projects)# запись в базу данных users_projekt
                            db.session.commit()# закрытие соеденения с бд
                return redirect(url_for('my_projects'))  
                
            # форма добавления фона
            if (listForm.get('theBackground')):
                file = request.files.getlist('file_bg')
                way = Config.linkPoto # указываю путь к папке
                for file in file:
                    if file.filename!="" or file.filename!=None:
                        print("я работаю")
                        file.save("app/" + os.path.join(way , file.filename))
                        #os.remove(way + file.filename)
                        patch = way + file.filename
                        prok_id = request.form.get('theBackground')
                        update = db.session.query(models.Project).filter_by(id = prok_id).one()# обновления таблицы Users_project
                        update.linkPoto = patch
                        db.session.commit()# создание соеденения
                        db.session.add(update)#обновление данных
                        db.session.commit()# закрытие соеденения с бд
                    else:
                        return redirect(url_for("my_projects"))
                return redirect(url_for('my_projects')) 
              
            
            # создание проекта
            if (listForm.get('projectName')):
                if (models.Project.query.filter_by(projectName = listForm['projectName']).first()):
                    error = "имя уже занято, проект не создался"
                    return render_template("my_projects.html", projects = projects, list = info, 
                        user_project = user_project, 
                        folders = folders,
                        tasks = tasks, error = error)
                pr_name = str(listForm["projectName"])
                functions.yadisk_check_pr(str(listForm["projectName"]) )# путь
                createProjekt = models.Project(projectName = pr_name,# создание проекта
                descProject = listForm['descProject'], linkDisk = Config.yad_pr_way+ pr_name, linkPoto = 'static/photoProjekt/phoenix-proj.png')
    
                db.session.commit()# создание соеденения
                db.session.add(createProjekt)# запись в базу данных users_projekt
                id_project = models.Project.query.filter_by(projectName = pr_name).first()
                print(id_project.id)
                for k, v in listForm.items():
                    if v == 'id_sel':
                        update_user_projects = models.Users_Projects(User_id = k, Project_id = id_project.id)# обновления таблицы Users_project
                        db.session.add(update_user_projects)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
                return redirect(url_for('my_projects'))
        return render_template("my_projects.html", projects = projects, list = info, 
                        user_project = user_project,error = error)
    if current_user.who == 2 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))

@application.route('/projects/<proj>/<topFolder>/', methods=['GET', 'POST'])  # Страница с папками и задачами шоуранера
@login_required  # только зарегистрированный человек сможет зайти
def projekt(proj,topFolder):
    if current_user.who==1 or current_user.who==32:
        #print(proj)# id проекта
        #print(topFolder)# верхняя папка
        project = models.Project.query.filter_by(id = proj).one()
        # ищем все папки и и все задачи этого проекта
        tasks = models.Tasks.query.filter_by(idProject = proj,id_folder = topFolder).all()# таблица задачь
        folders = models.Folder.query.filter_by(id_progect = proj,id_folder = topFolder).all()# таблица папок
        user_project = models.Users_Projects.query.filter_by(Project_id = proj).all()#запрос в базу данны для вывода какой человек к каому проекту
        massege = models.Massege.query.filter_by(id_progect = proj,id_folder = topFolder).all()# таблица сообщений
        user = models.Users_Data.query.all()
        
        if (topFolder == "-1"):
            print("")
            foldersTop = 0
            url = project.projectName
        else:
            foldersTop = models.Folder.query.filter_by(id_progect = proj,id = topFolder).one()# таблица верхней папки
            url = project.projectName + foldersTop.link
        print(folders)
        form = forms.CreateTask()# создание задачи
    
        if request.method == 'POST': # обработка post
            listForm = request.form.to_dict()# получение информации из формы
            print(listForm)
    
            # для создании задачи
            if (listForm.get('nameTask')):
                print("я создаю задачу")
                link = ""
                user = models.Users_Data.query.filter_by(idUser = listForm['id_sel']).one()
                if user.up_ooo == 0:
                    receipt = 0
                else:
                    receipt = 1
                if foldersTop == 0:
                    linkTask = project.projectName + "/"  + listForm['nameTask']  + "/" + user.name # вот тут  учитывается папка
                else:
                    linkTask = project.projectName + "/" + foldersTop.link + "/"  + listForm['nameTask']  + "/" + user.name # вот тут не учитывается папка
    
                link = functions.yadisk_check_pr(linkTask) # #путь, имя папки 
                print(link)
                createTask = models.Tasks(idUser = listForm['id_sel'], idProject = proj, id_folder = topFolder,
                nameTask = listForm['nameTask'],descTask = listForm['descTask'],timeTask = listForm['timeTask'],
                manyTask = listForm['manyTask'],statusСompleted = "Uncomplete",receipt=receipt,linkDisk = link)
                db.session.commit()# создание соеденения
                db.session.add(createTask)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд 
                return redirect(url_for('my_projects')) 
                        
            
            # для создании папки
            if (listForm.get('folderName')):
                print("я создаю папку")
                link =""
                if (topFolder == "-1"):
                    link = listForm['folderName']
                if (topFolder != "-1"):
                    link = foldersTop.link + "/" + listForm['folderName']
                createFolder = models.Folder(name = listForm['folderName'],id_progect = proj,link =link,
                    color="#c99f2d",id_folder=topFolder)
                db.session.commit()# создание соеденения
                db.session.add(createFolder)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
    
                functions.yadisk_check_pr(project.projectName + "/"+link)# путь
                return redirect(url_for('my_projects'))
    
            # для смены цвета папки
            if (listForm.get('colorTask')):
                print("Я изменяю цвет папки")
                print(listForm['colorTask'])
                print(listForm['color'])
                taskColor = models.Folder.query.filter_by(id = listForm['colorTask']).one()
                taskColor.color = listForm['color']
                db.session.commit()# создание соеденения
                db.session.add(taskColor)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
    
            # чат
            if (listForm.get('masseg')):
                print("записываю сообщение")
                user = models.Users_Data.query.filter_by(idUser = current_user.id).one() #User
                masg = models.Massege(id_progect = listForm['id_proj'],id_folder =listForm['id_folder'], 
                id_task = listForm['id_task'],id_user = listForm['id_user'], text = listForm['text'],img = user.avatar)
                db.session.commit()# создание соеденения
                db.session.add(masg)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
                massege = models.Massege.query.filter_by(id_progect = proj,id_folder = topFolder).all()# таблица сообщений
                get = {'success': 'true'}
                dic1 = {}
                for u in massege:
                    dic = u.__dict__
                    if '_sa_instance_state' in dic.keys():
                        del dic['_sa_instance_state']
                    name = 'mas'+ str(dic["id"])
                    dic1[name] = dic
                    get.update(dic1)
                #print(get)
                return json.dumps(get)
    
            # завершение задачи
            if (listForm.get('tascComplet')):
                task = models.Tasks.query.filter_by(id = listForm['tascComplet']).one()
                task.statusСompleted = "Сomplete"
                task.timeTask = date.today()
                db.session.commit()# создание соеденения
                db.session.add(task)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
    
            # переделать задачу
            if (listForm.get('tascChek')):
                task = models.Tasks.query.filter_by(id = listForm['tascChek']).one()
                task.statusСompleted = "Uncomplete"
                db.session.commit()# создание соеденения
                db.session.add(task)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд 
            # Редактировать задачу
            # if (listForm.get('task_edit')):
            #     task = models.Tasks.query.filter_by(id = listForm['task_edit']).one()
            #     if (request.form.get(edit_name)):
            #         print("я редактирую задачу")
            #         link = ""
            #         user = models.Users_Data.query.filter_by(idUser = task.idUser).one()
            #         if user.up_ooo == 0:
            #             receipt = 0
            #         else:
            #             receipt = 1
            #         if foldersTop == 0:
            #             linkTask = project.projectName + "/"  + request.form.get(edit_name)  + "/" + user.name # вот тут  учитывается папка
            #         else:
            #             linkTask = project.projectName + "/" + foldersTop.link + "/"  + request.form.get(edit_name)  + "/" + user.name # вот тут не учитывается папка
            #         link = functions.yadisk_check_pr(linkTask) # #путь, имя папки 
            #         print(link)
            #         editTask = models.Tasks(idUser = user.id, idProject = proj, id_folder = topFolder,
            #         nameTask = request.form.get(edit_name),descTask = request.form.get(edit_desc),timeTask = request.form.get(edit_time),
            #         manyTask = request.form.get(edit_money),statusСompleted = "Uncomplete",receipt=receipt,linkDisk = link)
            #         db.session.commit()# создание соеденения
            #         db.session.add(editTask)# запись в базу данных users_projekt
            #         db.session.commit()# закрытие соеденения с бд 
            #         return redirect(url_for('my_projects'))
            # Удалить задачу
            if (listForm.get('tascDelet')):
                task = models.Tasks.query.filter_by(id = listForm['tascDelet']).one()
                db.session.commit()# создание соеденения
                db.session.delete(task)# запись в базу данных users_projekt
                db.session.commit()# закрытие соеденения с бд
                return redirect(url_for('my_projects'))
        for u in tasks:
            yToken.publish(u.linkDisk)
            u.linkDisk = yToken.get_meta(u.linkDisk).public_url
        return render_template("projects.html",form=form,tasks = tasks,folders = folders,proj=proj,topFolder=topFolder,url=url,
                list = user_project,massege=massege,foldersTop=foldersTop,user=user)
    if current_user.who == 2 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))

        

@application.route('/salary', methods=['GET', 'POST'])  # Страница с зарплатами. Менеджер видит и устанавливает
@login_required  # только зарегистрированный человек сможет зайти
def salary():
    task = models.Tasks.query.filter_by(idUser = current_user.id).all()# поиск отмеченных пользователя
    total = 0
    for u in task:
        if u.statusСompleted=="Сomplete":
            total = int(u.manyTask) + total
    for u in task:
        u.timeTask = u.timeTask[8:10] + "." + u.timeTask[5:7] + "." + u.timeTask[0:4]
        u.timeTask = datetime.strptime(u.timeTask,'%d.%m.%Y').date()
        print(u.timeTask)
    a = date.today()# время сейчас
    print(a)
    b = date.today() - timedelta(days=7)#время неделю назад
    c = date.today() - timedelta(weeks = 4)#время месяц назад
    d = date.today() - timedelta(weeks = 52)#время год назад
    return render_template("salary.html", task = task, a = a, b = b, c = c, d = d, total = total)


@application.route('/employeers', methods=['GET', 'POST'])  # Страница с сотрудниками компании. Доступна менеджеру
@login_required  # только зарегистрированный человек сможет зайти
def employeers():
    if current_user.who==2 or current_user.who==32:
        info = models.User.query.all()
        if request.method == 'POST': # обработка post
            personal = request.form.get('person')# получение информации из формы
            user = models.User.query.filter_by(id = personal).one()
            #создаем договор
            # email  = models.User.query.filter_by(id = id).first()#поис email
            try:
                userData = user.pr.name + "\n" + "ИНН" + " " + user.pr.inn + "\t" + "ОГРН" + " " + user.pr.ogrn + "\n" + "Паспорт гражданина РФ" + "\n" + user.pr.passport + "\n" + "Выдан" + " " + user.pr.passportBy + " " + user.pr.passportData + "\n" "Код подразделения" + " " + user.pr.passportCod + "\n" + "Адрес:" + " " + user.pr.address + "\n" + "Номер счёта" + " " + user.pr.bankAccount + "\n" + "Банк получателя" + " " + user.pr.bankName + "\n" + " " + user.pr.bank_details + "\n" + user.email
            except TypeError:
                    user.pr.chekManager = 1
                    db.session.add(user)
                    db.session.commit()
                    return render_template('employeers.html', error="У пользователя " + user.pr.name + " не заполнен личный кабинет", list = info)
            context = { #шаблон для записи 
                    'date' : now.today().strftime("%d.%m.%Y"), #дата
                    'name':user.pr.name, #ФИО
                    'userData':userData, # реквезиты
                    }
            Config.doc.render(context)
            nameFail = "Договор " + user.pr.name + " " + now.today().strftime("%d.%m.%Y")#имя файла 
            Config.doc.save(Config.way_doc + nameFail + ".docx")#сохранение документа
            if yToken.exists(Config.yad_con + user.pr.name) == False:
                yToken.mkdir(Config.yad_con+ user.pr.name)# создание папки
                try:
                    yToken.upload(Config.way_doc+ nameFail +".docx", Config.yad_con + user.pr.name+"/" + nameFail +".docx")
                    os.remove(Config.way_doc+ nameFail +".docx")
                except yadisk.exceptions.PathExistsError:
                    user.pr.chekManager = 2
                    db.session.add(user)
                    db.session.commit()
                    os.remove(Config.way_doc+ nameFail +".docx")
                    return render_template('employeers.html',list=info, error="У " + user.pr.name + " Договор уже создан")
            elif yToken.exists(Config.yad_con+ user.pr.name) == True:#если папка существует
                try:
                    yToken.upload(Config.way_doc+ nameFail +".docx", Config.yad_con + user.pr.name+"/" + nameFail +".docx")
                    os.remove(Config.way_doc+ nameFail +".docx")
                except yadisk.exceptions.PathExistsError:
                    user.pr.chekManager = 2
                    db.session.add(user)
                    db.session.commit()
                    os.remove(Config.way_doc + nameFail +".docx")
                    return render_template('employeers.html',list=info, error="У " + user.pr.name + " Договор уже создан")
            user.pr.chekManager = 2
            db.session.add(user)
            db.session.commit()
            return redirect(url_for('employeers'))
        return render_template("employeers.html", list = info)
    if current_user.who == 1 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))


@application.route('/completed_tasks/<who>/', methods=['GET', 'POST'])  # Страница с выполненными задачами по людям
@login_required  # только зарегистрированный человек сможет зайти
def completed_tasks(who):
    if current_user.who==2 or current_user.who==32:
        user = models.User.query.filter_by(id = who).first()
        infoUser = models.Users_Data.query.filter_by(idUser = who).first()
        try:
            docList = list(yToken.listdir(Config.yad_con + infoUser.name))
            List = {}
            a = 0
            for i in docList:
                d = {docList[a].name : docList[a].file}
                List.update(d)
                a = a + 1  
        except yadisk.exceptions.PathNotFoundError:
            return render_template("completed_tasks.html",user = user,infoUser = infoUser,error = "Договоров нет")
    
        return render_template("completed_tasks.html",user = user,infoUser = infoUser, list = List)
    if current_user.who == 1 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))

@application.route('/cabinet_tasks_changer/<who>/', methods=['GET', 'POST'])  # Страница с выполненными задачами по людям
@login_required  # только зарегистрированный человек сможет зайти
def completed_tasks_changer(who):
    if current_user.who==2 or current_user.who==32:
        form = forms.PersonalForm()
        user = models.User.query.filter_by(id = who).first()
        infoUser = models.Users_Data.query.filter_by(idUser = who).first()
        if request.method == 'POST':
            listForm = request.form.to_dict()# получение информации из формы
            if form.name.data and str(form.name.data).strip() and form.name.data != None :
                user.name = form.name.data
                user.chekManager = 1
            if form.birthDAy.data and form.birthDAy.data != None :
                user.birthDAy = form.birthDAy.data
                user.chekManager = 1
            if form.passport.data and str(form.passport.data).strip() and form.passport.data != None :
                user.passport = form.passport.data
                user.chekManager = 1
            if form.passportData.data and form.passportData.data != None :
                user.passportData = form.passportData.data
                user.chekManager = 1
            if form.passportBy.data and form.passportBy.data != None :
                user.passportBy = form.passportBy.data
                user.chekManager = 1
            if form.passportCod.data and form.passportCod.data != None:
                user.passportCod = form.passportCod.data
                user.chekManager = 1
            if form.nickname.data and form.nickname.data != None :
                user.nickname = form.nickname.data
                user.chekManager = 1
            if form.link_vk.data and str(form.link_vk).strip() and form.link_vk.data != None :
                user.link_vk = form.link_vk.data
                user.chekManager = 1
            if form.inn.data and str(form.inn).strip() and form.inn.data != None:
                user.inn = form.inn.data
                user.chekManager = 1
            if form.bankAccount.data and str(form.bankAccount).strip() and form.bankAccount.data != None:
                user.bankAccount = form.bankAccount.data
                user.chekManager = 1
            if form.bank_details.data and str(form.bank_details).strip() and form.bank_details.data != None :
                user.bank_details = form.bank_details.data
                user.chekManager = 1
            if form.bankName.data and str(form.bankName).strip() and form.bankName.data != None :
                user.bankName = form.bankName.data
                user.chekManager = 1
            if form.phone_number.data and str(form.phone_number).strip() and form.phone_number.data != None:
                user.phone_number = form.phone_number.data
                user.chekManager = 1
            if form.address.data and str(form.address).strip() and form.address.data != None:
                user.address = form.address.data
                user.chekManager = 1
            if form.tags.data and str(form.tags).strip() and form.tags.data != None:
                user.tags = form.tags.data
                user.chekManager = 1
            if form.up_ooo.data != None :
                user.up_ooo = form.up_ooo.data
                user.chekManager = 1
            if form.ogrn.data and str(form.ogrn).strip() and form.ogrn.data != None :
                user.ogrn = form.ogrn.data
                user.chekManager = 1
            user.who=request.form.get("status")
            db.session.add(user)
            db.session.commit()
            return redirect(url_for('employeers'))
        return render_template("completed_tasks_changer.html",user = user,infoUser = infoUser,form=form)
    if current_user.who == 1 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))


@application.route('/create_contract', methods=['POST', 'GET'])
@login_required  # только зарегистрированный человек сможет зайти
def contract():
    if current_user.who==2 or current_user.who==32:
        info = models.User.query.all() #Получаем словарь с содержимым таблиц user и users_Data 
        if request.method == 'POST':
            listForm = request.form.to_dict()
            datStart = datetime.strptime(request.form['start'],'%Y-%m-%d')
            datEnd = datetime.strptime(request.form.get('end'),'%Y-%m-%d')
            random_day = functions.random_date(datStart,datEnd) # первая дата, последня дата / рандомная дата
            del listForm['start']#удаляем элемент времени из списка, что бы заработал цикл
            del listForm['end']#удаляем элемент времени из списка, что бы заработал цикл
            listFormKeys = listForm.keys()# выводим все ключи выбранных блоках
            # редактируем таблицу
            for id in listFormKeys:
                user = models.Users_Data.query.filter_by(idUser = id).one()#поиск юзера
                nameFail = "акт" + " " + user.name + " " + now.today().strftime("%d.%m.%Y")#имя файла 
                # вставляем дату
                context = { #шаблон для записи 
                    'date' : random_day.strftime("%d.%m.%Y"),#дата
                    'name': user.name, # имя пользователя
                }
                Config.doc_act.render(context)#ввод всех данных
                Config.doc_act.save(Config.way_doc + nameFail + ".docx")#сохранение документа 
                # вставка задач
                docx = Document(Config.way_doc + nameFail + ".docx")
                task = models.Tasks.query.filter_by(idUser = id).all()# поиск отмеченных пользователя
                i = 0
                for u in task:
                    if  datEnd >= datetime.strptime(u.timeTask,'%Y-%m-%d') >= datStart:
                        if u.statusСompleted == "Сomplete":
                            i = i + 1
                            if i > 1:
                                    docx.tables[1].add_row()
                            docx.tables[1]
                            docx.tables[1].cell(i,0).text = u.nameTask# 1 столбец, 2 строка 
                            docx.tables[1].cell(i,1).text = '50 дней с момента подписания настоящего соглашения'
                            docx.tables[1].cell(i,2).text = u.manyTask
                docx.save(Config.way_doc + nameFail +".docx")
                if yToken.exists(Config.yad_con + user.name + "/" ) == False:
                    yToken.mkdir(Config.yad_con + user.name + "/" )# создание папки
                    try:
                        yToken.upload(Config.way_doc+ nameFail +".docx", Config.yad_con + user.name + "/" + nameFail + ".docx")
                        os.remove(Config.way_doc + nameFail + ".docx")
                    except yadisk.exceptions.PathExistsError:
                        os.remove(Config.way_doc + nameFail + ".docx")
                        return render_template('employeers.html',list=info, error="У " + user.pr.name + " Договор уже создан")
                elif yToken.exists(Config.yad_con + user.name + "/") == True:#если папка существует
                    try:
                        yToken.upload(Config.way_doc+ nameFail +".docx",Config.yad_con + user.name + "/" + nameFail + ".docx")
                        os.remove(Config.way_doc + nameFail + ".docx")
                    except yadisk.exceptions.PathExistsError:
                        os.remove(Config.way_doc + nameFail + ".docx")
            return render_template('create_contract.html',list=info, error = "Документы созданны и добавлены на Яндекс диск и в личное дело")
        return render_template('create_contract.html', list=info)
    if current_user.who == 1 or current_user.who == 0 or current_user.who==3:
        return redirect(url_for('cabinet'))


@application.route('/admin', methods=['POST', 'GET'])  # Страница, доступная ЛИШЬ админу
@login_required  # только зарегистрированный человек сможет зайти
def admin():
    if current_user.id == 28:
        info = []
        info = models.User.query.all()
        tasks = models.Tasks.query.all()
        projects = models.Project.query.all()
        if request.method == "POST":
            user = models.Users_Data.query.filter_by(id = id).first()#поиск юзера
            task = models.Tasks.query.filter_by(idUser=id).all()
            user_project= models.Users_Projects.query.filter_by(User_id=id).all()
            db.session.delete(user)
            db.session.delete(task)
            db.session.delete(user_project)
            db.session.commit()
        return render_template('admin.html', list=info, tasks=tasks, projects=projects)
    else:
        return render_template('cabinet.html')

@application.route('/time',methods=['POST', 'GET'])# точное время jason
def get_time():
    print("на меня зашли")
    a = 0 # сколько задачь
    b = 0
    c = 0
    get = {'success': 'true'}
    data = {}
    if request.method == 'POST':
        listForm = request.form.to_dict()
        print(listForm)
        tasks = models.Tasks.query.filter_by(idUser=listForm['user'], statusСompleted = "Uncomplete").all()# таблица задачь
        for a in tasks:
            a = a + 1
        data = {'task': a}    
        get.update(data)
        
    return json.dumps(get)
    


# обработка ошибок
@application.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404


@application.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500
    # Обработчик ссылок

@application.route('/tutorial', methods=['GET', 'POST'])  #туториал
@login_required  # только зарегистрированный человек сможет зайти
def tutorial():
    return render_template('tutorial.html')

